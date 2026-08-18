import json
from pathlib import Path


def create_ml_analyzer_notebook():
    notebook_content = {
        "cells": [
            {
                "cell_type": "markdown",
                "metadata": {},
                "source": [
                    "# Stateful ML Analyzer Clustering\n",
                    "\n",
                    "## 1. Context & Overview\n",
                    "This notebook demonstrates how the **Smart AutoSorter AI Pro**'s core machine learning engine works under the hood. \n",
                    "The system utilizes an `IncrementalAnalyzer` that processes documents and automatically clusters them into semantic themes. This process allows developers and integration engineers to understand document themes without manual categorization.\n",
                    "\n",
                    "## 2. Parameter Explanations & Expectations\n",
                    "The `IncrementalAnalyzer` class requires several parameters to govern its behavior and state:\n",
                    "- `max_folders` (int): The upper limit of subdirectories the analyzer will generate for organizing documents (e.g., must be greater than 0 and is hard-capped at 12 by system constraints).\n",
                    "- `stop_words` (set): A set of common words (e.g., 'the', 'and', 'for') that are filtered out during tokenization and TF-IDF calculation to ensure meaningful clustering.\n",
                    "- `db` (Database): An active `Database` instance. The analyzer is stateful and stores TF-IDF vocabularies, document mappings, and other metadata inside this SQLite-backed database.\n",
                    "- `strategy_name` (str): Specifies the clustering strategy. The fully offline/deterministic strategy is `'default'`, which uses recursive KMeans and TF-IDF keywords. `'generative'` can be used when LLM-driven naming is active.\n",
                    "- `model_path` (str/None): The filesystem path to semantic embedding model weights if neural vector clustering is enabled. When `None` or omitted, it falls back to standard keyword-based modeling, which has zero network or hardware model-loading dependencies.\n",
                    "\n",
                    "### Expected Inputs:\n",
                    "- **Base Directory (`base_dir`)**: The root directory being analyzed.\n",
                    "- **Corpus (`new_corpus`)**: A dictionary mapping file relative paths to their extracted text content (or dictionaries with `'text'` and `'hash'` keys).\n",
                    "\n",
                    "### Expected Outputs:\n",
                    "- **Sorting Plan**: A nested dictionary representation of the proposed directory structure containing files mapped to destination paths and categories based on mathematical clustering profiles.",
                ],
            },
            {
                "cell_type": "code",
                "execution_count": None,
                "metadata": {},
                "outputs": [],
                "source": [
                    "import json\n",
                    "import tempfile\n",
                    "from pathlib import Path\n",
                    "\n",
                    "from app.core.analyzer import IncrementalAnalyzer\n",
                    "from app.core.db import Database\n",
                    "\n",
                    "# Core imports from Smart AutoSorter AI Pro\n",
                    "from app.core.db_worker import DBWorker",
                ],
            },
            {
                "cell_type": "markdown",
                "metadata": {},
                "source": [
                    "## 3. Sandboxing & Safe Database Initialization\n",
                    "To prevent modifying any live user data, we set up a temporary directory to serve as our isolated sandbox. All document processing and database writes will occur within this folder.",
                ],
            },
            {
                "cell_type": "code",
                "execution_count": None,
                "metadata": {},
                "outputs": [],
                "source": [
                    "# Create an isolated sandbox environment\n",
                    "sandbox_dir = tempfile.TemporaryDirectory()\n",
                    "base_dir = sandbox_dir.name\n",
                    'print(f"[*] Sandbox directory initialized safely at: {base_dir}")\n',
                    "\n",
                    "# Instantiate the background database worker and local SQLite DB\n",
                    "db_worker = DBWorker()\n",
                    'db_path = Path(base_dir) / "sandbox_autosorter.db"\n',
                    "db = Database(db_path, db_worker)\n",
                    'print(f"[+] Sandbox Database initialized successfully at: {db_path}")',
                ],
            },
            {
                "cell_type": "markdown",
                "metadata": {},
                "source": [
                    "## 4. Analyzer Initialization\n",
                    "Now, we initialize our stateful `IncrementalAnalyzer`. We'll configure it with a maximum limit of 3 folders, a clean set of stop words, and use the `'default'` KMeans clustering strategy.",
                ],
            },
            {
                "cell_type": "code",
                "execution_count": None,
                "metadata": {},
                "outputs": [],
                "source": [
                    "# Stop words to filter out noise words during clustering\n",
                    'stop_words = {"the", "and", "for", "this", "that", "with", "from", "your", "will", "are", "not", "can"}\n',
                    "\n",
                    "analyzer = IncrementalAnalyzer(\n",
                    "    max_folders=3,\n",
                    "    stop_words=stop_words,\n",
                    "    db=db,\n",
                    '    strategy_name="default",\n',
                    "    model_path=None  # Fallback to local keyword-based recursive KMeans strategy\n",
                    ")\n",
                    'print("[+] Stateful IncrementalAnalyzer successfully initialized!")',
                ],
            },
            {
                "cell_type": "markdown",
                "metadata": {},
                "source": [
                    "## 5. Feeding Documents (Incremental Training)\n",
                    "We define a diverse sample corpus spanning three distinct themes: **Finance**, **Technology**, and **Healthcare**. We train the analyzer incrementally by feeding chunks using `partial_fit`.",
                ],
            },
            {
                "cell_type": "code",
                "execution_count": None,
                "metadata": {},
                "outputs": [],
                "source": [
                    "sample_corpus = {\n",
                    '    "invoice_1042.txt": "invoice billing statement payment wire transfer balance sheet finance department banking profit revenue expense ledger audit audit",\n',
                    '    "quarterly_report.txt": "finance banking revenue quarterly profits balance sheet asset liability stock market investment ledger statement account billing payment",\n',
                    '    "neural_net_notes.txt": "machine learning artificial intelligence deep learning algorithms computer science software python neural networks model train GPU CPU programming git",\n',
                    '    "api_integration.txt": "software engineering python computer science developers api git source code debug program system architecture database query server",\n',
                    '    "clinical_trial_a.txt": "medical patient medicine health clinical trial cardiology pharmaceutical dosage diagnosis therapy disease doctor physician hospital treatment",\n',
                    '    "patient_health_summary.txt": "health patient medical clinic diagnosis doctor medicine therapy hospital pharmaceutical cardiology trial disease treatment blood pressure dosage"\n',
                    "}\n",
                    "\n",
                    'print("[*] Training the stateful analyzer on the sample corpus...")\n',
                    "analyzer.partial_fit(base_dir, sample_corpus)\n",
                    'print("[+] Partial fit complete. All documents ingested and TF-IDF tables populated.")',
                ],
            },
            {
                "cell_type": "markdown",
                "metadata": {},
                "source": [
                    "## 6. Generating the Sorting Plan\n",
                    "With the analyzer trained, we generate a sorting plan mapping files to proposed directory folders. The analyzer automatically inspects the themes of our files and clusters them accordingly.",
                ],
            },
            {
                "cell_type": "code",
                "execution_count": None,
                "metadata": {},
                "outputs": [],
                "source": [
                    'print("[*] Generating sorting plan based on semantic similarities...")\n',
                    "plan = analyzer.generate_sorting_plan(base_dir)\n",
                    "\n",
                    'print("\\n[+] Proposed Sorting Plan:")\n',
                    "print(json.dumps(plan, indent=2))",
                ],
            },
            {
                "cell_type": "markdown",
                "metadata": {},
                "source": [
                    "## 7. Clean Resource Teardown\n",
                    "Lastly, we ensure all background workers and temporary resources are terminated and cleaned up properly.",
                ],
            },
            {
                "cell_type": "code",
                "execution_count": None,
                "metadata": {},
                "outputs": [],
                "source": [
                    'print("[*] Terminating analyzer and database background threads...")\n',
                    "analyzer.terminate()\n",
                    "db_worker.stop()\n",
                    "sandbox_dir.cleanup()\n",
                    'print("[+] Sandbox environment cleaned up successfully. Bye!")',
                ],
            },
        ],
        "metadata": {
            "kernelspec": {
                "display_name": "Python 3 (ipykernel)",
                "language": "python",
                "name": "python3",
            },
            "language_info": {"name": "python"},
        },
        "nbformat": 4,
        "nbformat_minor": 2,
    }
    return notebook_content


def create_multi_format_text_extraction_notebook():
    notebook_content = {
        "cells": [
            {
                "cell_type": "markdown",
                "metadata": {},
                "source": [
                    "# Multi-Format Text Extraction & Local Session Management\n",
                    "\n",
                    "## 1. Context & Overview\n",
                    "This notebook demonstrates the document ingestion layer of **Smart AutoSorter AI Pro**. \n",
                    "We will explore:\n",
                    "1. **Local Session Management**: Initializing and managing the lifespan of an application run via `AppSession`.\n",
                    "2. **Multi-Format Text Extraction**: Writing programmatically supported formats (`.txt`, `.csv`, `.docx`, `.xlsx`) to disk and extracting their text content safely.\n",
                    "3. **Database Inspection**: Inspecting the internal SQLite databases to verify that file text, hashes, and schemas are correctly registered in the system's storage layer.",
                ],
            },
            {
                "cell_type": "code",
                "execution_count": None,
                "metadata": {},
                "outputs": [],
                "source": [
                    "import csv\n",
                    "import os\n",
                    "import tempfile\n",
                    "\n",
                    "# Install helpers/libraries for office formats\n",
                    "import docx\n",
                    "import openpyxl\n",
                    "\n",
                    "# Core Smart AutoSorter imports\n",
                    "from app.config import Settings\n",
                    "from app.core.db_conn import get_db_connection\n",
                    "from app.core.extractor import build_corpus_generator\n",
                    "from app.core.session import AppSession",
                ],
            },
            {
                "cell_type": "markdown",
                "metadata": {},
                "source": [
                    "## 2. Sandbox Setup & Generating Sample Documents\n",
                    "To guarantee isolated execution, we create a temporary directory and generate files across various supported formats.",
                ],
            },
            {
                "cell_type": "code",
                "execution_count": None,
                "metadata": {},
                "outputs": [],
                "source": [
                    "# Set up a secure, isolated temporary workspace\n",
                    "sandbox_dir = tempfile.TemporaryDirectory()\n",
                    "base_dir = sandbox_dir.name\n",
                    'print(f"[*] Safe base workspace directory: {base_dir}")\n',
                    "\n",
                    "# 1. Create a plain text (.txt) file\n",
                    'txt_path = os.path.join(base_dir, "sample_notes.txt")\n',
                    'with open(txt_path, "w", encoding="utf-8") as f:\n',
                    '    f.write("Software development notes. Python programming, unit tests, and continuous integration pipelines.")\n',
                    "\n",
                    "# 2. Create a CSV (.csv) spreadsheet file\n",
                    'csv_path = os.path.join(base_dir, "sample_spreadsheet.csv")\n',
                    'with open(csv_path, "w", encoding="utf-8", newline="") as f:\n',
                    "    writer = csv.writer(f)\n",
                    '    writer.writerow(["Item", "Amount", "Department"])\n',
                    '    writer.writerow(["Servers", "25000", "Infrastructure"])\n',
                    '    writer.writerow(["Monitors", "12000", "Hardware"])\n',
                    "\n",
                    "# 3. Create a Microsoft Word (.docx) file\n",
                    'docx_path = os.path.join(base_dir, "sample_doc.docx")\n',
                    "doc = docx.Document()\n",
                    'doc.add_heading("Project Specification", level=1)\n',
                    'doc.add_paragraph("This document details clinical trial results, healthcare diagnostics, and medical treatments.")\n',
                    "doc.save(docx_path)\n",
                    "\n",
                    "# 4. Create an Excel (.xlsx) workbook file\n",
                    'xlsx_path = os.path.join(base_dir, "sample_ledger.xlsx")\n',
                    "wb = openpyxl.Workbook()\n",
                    "ws = wb.active\n",
                    'ws.title = "Financial Overview"\n',
                    'ws["A1"] = "Quarterly Budget Report"\n',
                    'ws["A2"] = "Total Investment Assets and Balance Sheets"\n',
                    'ws["B2"] = 750000\n',
                    "wb.save(xlsx_path)\n",
                    "\n",
                    "files_created = [os.path.basename(p) for p in [txt_path, csv_path, docx_path, xlsx_path]]\n",
                    'print(f"[+] Generated {len(files_created)} sample multi-format files: {files_created}")',
                ],
            },
            {
                "cell_type": "markdown",
                "metadata": {},
                "source": [
                    "## 3. Session Initialization & Core Ingestion Pipeline\n",
                    "We now instantiate our custom `AppSession` utilizing explicit `Settings`. We then trigger the synchronous `build_corpus_generator` to run text extraction across all the generated formats.",
                ],
            },
            {
                "cell_type": "code",
                "execution_count": None,
                "metadata": {},
                "outputs": [],
                "source": [
                    "# Initialize settings with AI consent disabled for deterministic, lightweight extraction\n",
                    "settings = Settings(AI_CONSENT_GRANTED=False, MAX_FOLDERS=5)\n",
                    "\n",
                    "# Initialize AppSession\n",
                    "session = AppSession(settings=settings, base_dir=base_dir)\n",
                    'print(f"[+] AppSession initialized! ID: {session.session_id}")\n',
                    'print(f"[*] Session directory containing logs & databases: {session.session_dir}")\n',
                    "\n",
                    "def progress_callback():\n",
                    '    """Progress callback function for text extraction."""\n',
                    '    print("    -> Extracting file...")\n',
                    "\n",
                    "# Execute synchronous extraction of text payload and hashing from sample documents\n",
                    "generator = build_corpus_generator(\n",
                    "    base_dir=base_dir,\n",
                    "    items_to_sort=files_created,\n",
                    "    progress_callback=progress_callback,\n",
                    "    max_workers=2,\n",
                    "    db=session.db,\n",
                    "    chunk_size=10,\n",
                    "    sequential=True,\n",
                    "    settings=settings\n",
                    ")\n",
                    "\n",
                    "extracted_chunks = list(generator)\n",
                    'print(f"\\n[+] Ingestion complete! Processed {len(extracted_chunks)} chunks.")',
                ],
            },
            {
                "cell_type": "markdown",
                "metadata": {},
                "source": [
                    "## 4. Database Inspection via SQL\n",
                    "We connect directly to the session's SQLite database file (`autosorter.db`) to inspect the tables. This allows developers to audit the exact extracted text strings and calculated unique file hashes stored persistently.",
                ],
            },
            {
                "cell_type": "code",
                "execution_count": None,
                "metadata": {},
                "outputs": [],
                "source": [
                    'db_file = session.session_dir / "autosorter.db"\n',
                    'print(f"[*] Connecting to session database: {db_file}")\n',
                    "\n",
                    "conn = get_db_connection(str(db_file))\n",
                    "cursor = conn.cursor()\n",
                    "\n",
                    "# 1. Fetch tables schema list\n",
                    "cursor.execute(\"SELECT name FROM sqlite_master WHERE type='table';\")\n",
                    "tables = cursor.fetchall()\n",
                    'print(f"[+] Database Tables present: {[t[0] for t in tables]}")\n',
                    "\n",
                    "# 2. Query documents table\n",
                    'cursor.execute("SELECT filepath, file_hash, extracted_text FROM documents;")\n',
                    "records = cursor.fetchall()\n",
                    "\n",
                    'print("\\n--- Persisted Documents & Extracted Content ---")\n',
                    "for filepath, f_hash, text in records:\n",
                    '    truncated_text = text[:100] + "..." if len(text) > 100 else text\n',
                    '    print(f"File: {filepath}")\n',
                    '    print(f"  Hash: {f_hash}")\n',
                    '    print(f"  Text preview: {truncated_text}")\n',
                    '    print("-" * 45)\n',
                    "\n",
                    "conn.close()",
                ],
            },
            {
                "cell_type": "markdown",
                "metadata": {},
                "source": [
                    "## 5. Clean Resource Teardown\n",
                    "Finally, close the session to release file locks on databases, and safely delete the temporary sandbox.",
                ],
            },
            {
                "cell_type": "code",
                "execution_count": None,
                "metadata": {},
                "outputs": [],
                "source": [
                    'print("[*] Closing session and cleaning up directories...")\n',
                    "session.close()\n",
                    "sandbox_dir.cleanup()\n",
                    'print("[+] Environment cleaned up. Extraction workflow complete!")',
                ],
            },
        ],
        "metadata": {
            "kernelspec": {
                "display_name": "Python 3 (ipykernel)",
                "language": "python",
                "name": "python3",
            },
            "language_info": {"name": "python"},
        },
        "nbformat": 4,
        "nbformat_minor": 2,
    }
    return notebook_content


def create_virtual_sorting_verification_notebook():
    notebook_content = {
        "cells": [
            {
                "cell_type": "markdown",
                "metadata": {},
                "source": [
                    "# Virtual Sorting Verification & Safe Simulation Testing\n",
                    "\n",
                    "## 1. Context & Overview\n",
                    "Before performing physical move/rename operations on a user's filesystem, **Smart AutoSorter AI Pro** executes a proactive, in-memory dry run verification using the `VerificationEngine`.\n",
                    "\n",
                    "This safety check guarantees:\n",
                    "1. **No dynamic path collisions**: Preventing multiple source files from overwriting each other if they are assigned to the same target path.\n",
                    "2. **No invalid/long paths**: Detecting and warning if any generated destination paths violate the standard operating system character limits (e.g. 260 character limit on Windows).\n",
                    "3. **Safe Simulation**: Performing validation entirely in memory without actually writing, copying, or deleting any files on disk.",
                ],
            },
            {
                "cell_type": "code",
                "execution_count": None,
                "metadata": {},
                "outputs": [],
                "source": [
                    "import json\n",
                    "import os\n",
                    "import tempfile\n",
                    "\n",
                    "# Core Smart AutoSorter imports\n",
                    "from app.core.verifier import VerificationEngine",
                ],
            },
            {
                "cell_type": "markdown",
                "metadata": {},
                "source": [
                    "## 2. Sandbox Setup\n",
                    "We'll instantiate a clean sandbox directory to mock our local filesystem base path.",
                ],
            },
            {
                "cell_type": "code",
                "execution_count": None,
                "metadata": {},
                "outputs": [],
                "source": [
                    "# Setup mock workspace\n",
                    "sandbox_dir = tempfile.TemporaryDirectory()\n",
                    "base_dir = os.path.normpath(sandbox_dir.name)\n",
                    'print(f"[*] Safe mock workspace base dir: {base_dir}")',
                ],
            },
            {
                "cell_type": "markdown",
                "metadata": {},
                "source": [
                    "## 3. Scenario 1: Normal Sorting Operations (Successful Validation)\n",
                    "We configure a clean, structured sorting plan with non-conflicting destinations. The `VerificationEngine` should run a simulation and return a successful verification status.",
                ],
            },
            {
                "cell_type": "code",
                "execution_count": None,
                "metadata": {},
                "outputs": [],
                "source": [
                    "normal_plan = {\n",
                    '    "document_A.txt": {\n',
                    '        "__type__": "file",\n',
                    '        "relative_source": "document_A.txt",\n',
                    '        "target_filename": "document_A.txt",\n',
                    "    },\n",
                    '    "Finance": {\n',
                    '        "__type__": "directory",\n',
                    '        "document_B.txt": {\n',
                    '            "__type__": "file",\n',
                    '            "relative_source": "document_B.txt",\n',
                    '            "target_filename": "renamed_B.txt",\n',
                    "        }\n",
                    "    }\n",
                    "}\n",
                    "\n",
                    'print("[*] Verifying normal sorting plan integrity...")\n',
                    "result = VerificationEngine.verify_plan_integrity(base_dir, normal_plan)\n",
                    "\n",
                    'print("\\n[+] Verification Result:")\n',
                    "print(f\"  Success: {result['success']}\")\n",
                    "print(f\"  Warnings: {result['warnings']}\")\n",
                    "print(f\"  Collisions: {result['collisions']}\")",
                ],
            },
            {
                "cell_type": "markdown",
                "metadata": {},
                "source": [
                    "## 4. Scenario 2: Dynamic Path Collision Detection\n",
                    "Now, let's intentionally introduce a collision where two separate source files are directed to the same destination target file. The `VerificationEngine` must proactively detect this and mark the plan verification as unsuccessful.",
                ],
            },
            {
                "cell_type": "code",
                "execution_count": None,
                "metadata": {},
                "outputs": [],
                "source": [
                    "colliding_plan = {\n",
                    '    "Finance": {\n',
                    '        "__type__": "directory",\n',
                    '        "invoice_1.txt": {\n',
                    '            "__type__": "file",\n',
                    '            "relative_source": "invoice_1.txt",\n',
                    '            "target_filename": "clashing_name.txt",\n',
                    "        },\n",
                    '        "receipt_9.txt": {\n',
                    '            "__type__": "file",\n',
                    '            "relative_source": "receipt_9.txt",\n',
                    '            "target_filename": "clashing_name.txt", # Collision!\n',
                    "        }\n",
                    "    }\n",
                    "}\n",
                    "\n",
                    'print("[*] Verifying colliding plan integrity...")\n',
                    "result = VerificationEngine.verify_plan_integrity(base_dir, colliding_plan)\n",
                    "\n",
                    'print("\\n[-] Verification Result:")\n',
                    "print(f\"  Success: {result['success']}\")\n",
                    "print(f\"  Collisions: {json.dumps(result['collisions'], indent=2)}\")\n",
                    "print(f\"  Warnings: {result['warnings']}\")",
                ],
            },
            {
                "cell_type": "markdown",
                "metadata": {},
                "source": [
                    "## 5. Scenario 3: Windows Path Length Warnings Limit Detection\n",
                    "The system warns about paths exceeding the standard limit. We'll simulate a plan containing a filename with 250 characters.",
                ],
            },
            {
                "cell_type": "code",
                "execution_count": None,
                "metadata": {},
                "outputs": [],
                "source": [
                    'long_filename = "x" * 250 + ".txt"\n',
                    "long_path_plan = {\n",
                    '    "short_file.txt": {\n',
                    '        "__type__": "file",\n',
                    '        "relative_source": "short_file.txt",\n',
                    '        "target_filename": long_filename,\n',
                    "    }\n",
                    "}\n",
                    "\n",
                    'print("[*] Verifying long-path plan integrity...")\n',
                    "result = VerificationEngine.verify_plan_integrity(base_dir, long_path_plan)\n",
                    "\n",
                    'print("\\n[-] Verification Result:")\n',
                    "print(f\"  Success: {result['success']}\")\n",
                    "print(f\"  Long Paths Found: {len(result['long_paths'])}\")\n",
                    "if result['long_paths']:\n",
                    "    print(f\"  Example long path truncated: {result['long_paths'][0]['path'][:80]}...\")\n",
                    "print(f\"  Warnings: {result['warnings']}\")",
                ],
            },
            {
                "cell_type": "markdown",
                "metadata": {},
                "source": [
                    "## 6. Safe Workspace Cleanup\n",
                    "Clean up the temporary sandbox workspace safely.",
                ],
            },
            {
                "cell_type": "code",
                "execution_count": None,
                "metadata": {},
                "outputs": [],
                "source": [
                    'print("[*] Cleaning up mock sandbox base directory...")\n',
                    "sandbox_dir.cleanup()\n",
                    'print("[+] Workspace deleted. Verification simulation finalized safely!")',
                ],
            },
        ],
        "metadata": {
            "kernelspec": {
                "display_name": "Python 3 (ipykernel)",
                "language": "python",
                "name": "python3",
            },
            "language_info": {"name": "python"},
        },
        "nbformat": 4,
        "nbformat_minor": 2,
    }
    return notebook_content


def main():
    notebooks_dir = Path("notebooks")
    notebooks_dir.mkdir(parents=True, exist_ok=True)

    # 1. Stateful ML analyzer clustering
    ml_nb = create_ml_analyzer_notebook()
    with open(
        notebooks_dir / "01_ml_analyzer_clustering.ipynb", "w", encoding="utf-8"
    ) as f:
        json.dump(ml_nb, f, indent=1)
    print("[+] Generated 01_ml_analyzer_clustering.ipynb")

    # 2. Multi-format text extraction & session management
    ex_nb = create_multi_format_text_extraction_notebook()
    with open(
        notebooks_dir / "02_multi_format_text_extraction.ipynb", "w", encoding="utf-8"
    ) as f:
        json.dump(ex_nb, f, indent=1)
    print("[+] Generated 02_multi_format_text_extraction.ipynb")

    # 3. Virtual sorting verification
    ver_nb = create_virtual_sorting_verification_notebook()
    with open(
        notebooks_dir / "03_virtual_sorting_verification.ipynb", "w", encoding="utf-8"
    ) as f:
        json.dump(ver_nb, f, indent=1)
    print("[+] Generated 03_virtual_sorting_verification.ipynb")


if __name__ == "__main__":
    main()
